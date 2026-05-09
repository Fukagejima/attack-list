"""
Word (.docx) マスキング処理
- 段落・ラン・テーブル・ヘッダー/フッター・ハイパーリンクをすべて処理
- レイアウト・スタイルは一切変更しない
"""
from __future__ import annotations
import io
import copy

from docx import Document
from docx.oxml.ns import qn

from .patterns import apply_masks


def _mask_run_text(run, opts: dict) -> int:
    """1つのランのテキストをマスクして変更数を返す"""
    if not run.text:
        return 0
    masked, cnt = apply_masks(run.text, opts)
    if cnt > 0:
        run.text = masked
    return cnt


def _mask_paragraph(para, opts: dict) -> int:
    """段落内の全ランをマスクして合計変更数を返す"""
    total = 0
    for run in para.runs:
        total += _mask_run_text(run, opts)
    return total


def _mask_hyperlinks(doc: Document, opts: dict) -> int:
    """
    ドキュメント内のハイパーリンクURLをマスク
    - 関係性ファイル内のターゲットURL
    - 段落内の w:hyperlink 表示テキスト（ランとして処理済みだが念のため）
    """
    if not opts.get('url'):
        return 0

    count = 0
    for rel in doc.part.rels.values():
        if 'hyperlink' in rel.reltype.lower():
            original = rel._target
            masked, cnt = apply_masks(original, opts)
            if cnt > 0:
                rel._target = masked
                count += cnt

    # インラインハイパーリンクのターゲット属性も処理
    for elem in doc.element.iter(qn('w:hyperlink')):
        anchor = elem.get(qn('w:anchor'))
        if anchor:
            masked, cnt = apply_masks(anchor, opts)
            if cnt > 0:
                elem.set(qn('w:anchor'), masked)
                count += cnt

    return count


def _mask_text_boxes(doc: Document, opts: dict) -> int:
    """テキストボックス内のテキストをマスク"""
    count = 0
    for txbx in doc.element.iter(qn('w:txbxContent')):
        for p_elem in txbx.iter(qn('w:p')):
            for r_elem in p_elem.iter(qn('w:r')):
                t_elem = r_elem.find(qn('w:t'))
                if t_elem is not None and t_elem.text:
                    masked, cnt = apply_masks(t_elem.text, opts)
                    if cnt > 0:
                        t_elem.text = masked
                        count += cnt
    return count


def mask_word(file_bytes: bytes, opts: dict) -> tuple[bytes, int]:
    """
    Word ファイルをマスクして返す。
    return: (masked_bytes, total_mask_count)
    """
    doc = Document(io.BytesIO(file_bytes))
    total = 0

    # ── 本文段落 ──
    for para in doc.paragraphs:
        total += _mask_paragraph(para, opts)

    # ── テーブル ──
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total += _mask_paragraph(para, opts)

    # ── ヘッダー / フッター ──
    for section in doc.sections:
        for hdr in (section.header, section.footer,
                    section.even_page_header, section.even_page_footer,
                    section.first_page_header, section.first_page_footer):
            for para in hdr.paragraphs:
                total += _mask_paragraph(para, opts)
            for table in hdr.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            total += _mask_paragraph(para, opts)

    # ── ハイパーリンク URL ──
    total += _mask_hyperlinks(doc, opts)

    # ── テキストボックス ──
    total += _mask_text_boxes(doc, opts)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), total
